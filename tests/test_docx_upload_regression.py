"""回归测试：docx/xlsx 必须能走通线上的真实解析路径。

背景（Phase 0 发现的 bug）：ALLOWED_EXTENSIONS 允许上传 .docx/.xlsx，但线上
insert_into_index -> utils/llama.py:get_nodes_from_file -> SimpleDirectoryReader
解析这两类文件分别依赖 docx2txt / openpyxl。这两个包曾经不在主依赖里，导致
上传 docx/xlsx 时解析直接抛 ImportError（被 upload 接口的 except 吞掉，用户
只看到"文件处理出错"）。

这组测试用真实的最小 docx/xlsx 文件（python-docx / openpyxl 现场生成，不 mock
任何解析器）走 get_nodes_from_file，断言能解析出内容——如果依赖再丢，这里会
立刻红。风格参考 tests/test_document_parsers.py（那边测的是 utils/file.py 的
read_file_contents 路径，这里测的是索引插入用的 SimpleDirectoryReader 路径，
两条路径互不覆盖）。
"""
import os
import unittest

from utils.llama import get_nodes_from_file

import tests._pathsetup  # noqa: F401  (adds backend/app to sys.path)


class DocxUploadRegressionTest(unittest.TestCase):
    def test_docx_parses_through_production_ingest_path(self):
        """get_nodes_from_file 对真实 .docx 必须解析出文本节点，而不是 ImportError。"""
        import tempfile

        from docx import Document as DocxDocument

        with tempfile.TemporaryDirectory() as tmp_dir:
            docx_path = os.path.join(tmp_dir, "回归测试文档.docx")
            doc = DocxDocument()
            doc.add_paragraph("成都信息工程大学的校训是成于大气 信达天下。")
            doc.add_paragraph("这是 docx 上传解析路径的回归测试内容。")
            doc.save(docx_path)

            nodes = get_nodes_from_file(docx_path)

        self.assertGreater(len(nodes), 0, "docx 解析不出任何节点")
        all_text = "".join(node.get_content() for node in nodes)
        self.assertIn("成于大气", all_text)
        self.assertIn("回归测试内容", all_text)
        # 文件名进 metadata（评测和去重都靠它），doc id 以文件名开头
        # （filename_as_id=True，DocxReader 会追加 _part_N 后缀）
        self.assertEqual(nodes[0].metadata.get("file_name"), "回归测试文档.docx")
        self.assertTrue(str(nodes[0].ref_doc_id).startswith("回归测试文档.docx"))

    def test_xlsx_parses_through_production_ingest_path(self):
        """get_nodes_from_file 对真实 .xlsx 必须解析出内容，而不是 ImportError。"""
        import tempfile

        from openpyxl import Workbook

        with tempfile.TemporaryDirectory() as tmp_dir:
            xlsx_path = os.path.join(tmp_dir, "回归测试表格.xlsx")
            wb = Workbook()
            ws = wb.active
            ws.append(["部门", "电话"])
            ws.append(["一卡通服务中心", "028-85966666"])
            wb.save(xlsx_path)

            nodes = get_nodes_from_file(xlsx_path)

        self.assertGreater(len(nodes), 0, "xlsx 解析不出任何节点")
        all_text = "".join(node.get_content() for node in nodes)
        self.assertIn("一卡通服务中心", all_text)


if __name__ == "__main__":
    unittest.main()
