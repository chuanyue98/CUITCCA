"""上传/QA 生成链路（``utils/file.py.read_file_contents``）的解析行为测试。

这个文件原本直接 mock ``utils.file`` 内部的 ``tempfile``/``pdfplumber``——那
是在 ``_read_file_sync`` 自己实现 docx/pdf/xlsx 解析分支的年代写的。解析实现
已经统一收敛到 ``handlers/parsers`` 注册表（两条链路共用一套），那些针对内部
实现细节的 mock 随之失效。这里改成按**行为契约**测试：给 ``read_file_contents``
一个 UploadFile，断言拿到的文本，不再假设它内部用什么库、要不要落临时文件。

注册表本身和各格式解析器的单元测试在 ``test_parsers_registry.py``。
"""
import asyncio
import io
import unittest
from io import BytesIO
from unittest.mock import MagicMock, patch

from docx import Document as DocxDocument
from fastapi import UploadFile
from handlers.parsers.types import DocumentParseError, ParserUnavailableError
from utils.file import read_file_contents

import tests._pathsetup  # noqa: F401


def _docx_bytes() -> bytes:
    """造一个真实的 .docx（段落 + 表格）。

    比 mock ``docx.Document`` 更可靠：新的 docx 解析器要按文档原始顺序遍历
    ``document.element.body`` 的子元素来混合抽取段落和表格，mock 出来的假对象
    没有真实 XML 结构，测不出这个行为，反而会因为 mock 的实现细节漂移而误报。
    """
    doc = DocxDocument()
    doc.add_paragraph("Hello World.")
    doc.add_paragraph("This is a docx test.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "姓名"
    table.cell(0, 1).text = "电话"
    table.cell(1, 0).text = "段老师"
    table.cell(1, 1).text = "028-85967803"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class DocumentParsersTest(unittest.TestCase):
    def test_docx_paragraphs_are_extracted(self):
        upload = UploadFile(filename="test.docx", file=BytesIO(_docx_bytes()))
        content = asyncio.run(read_file_contents(upload))
        self.assertIn("Hello World.", content)
        self.assertIn("This is a docx test.", content)

    def test_docx_table_content_is_extracted(self):
        """旧实现只读 ``doc.paragraphs``，表格内容整个丢失。这是本次改造修掉的
        真实缺陷——语料里大量 docx 是申请表/联系方式表，正文几乎全在表格里。"""
        upload = UploadFile(filename="test.docx", file=BytesIO(_docx_bytes()))
        content = asyncio.run(read_file_contents(upload))
        self.assertIn("段老师", content)
        self.assertIn("028-85967803", content)

    @patch("pdfplumber.open")
    def test_pdf_pages_are_joined_with_separator(self, mock_pdf_open):
        """页与页之间要有分隔。

        旧实现是 ``content = content + page_text`` 直接拼接，上一页末尾的字和
        下一页开头的字会粘成一个词（"第1页。第2页" -> "第1页。第2页" 中间无
        空白），既影响可读性也影响分词/检索。现在按页用空行连接，经上传链路
        的空白归一化后表现为一个空格。
        """
        mock_pdf = MagicMock()
        page1, page2 = MagicMock(), MagicMock()
        page1.extract_text.return_value = "Hello PDF page 1."
        page2.extract_text.return_value = "Hello PDF page 2."
        # find_tables() 返回空 -> 走"整页纯文本"分支（MagicMock 的 __iter__
        # 默认就是空迭代器，这里显式写出来让意图明确，不依赖 mock 的默认行为）
        page1.find_tables.return_value = []
        page2.find_tables.return_value = []
        mock_pdf.pages = [page1, page2]
        mock_pdf_open.return_value.__enter__.return_value = mock_pdf

        upload = UploadFile(filename="test.pdf", file=BytesIO(b"fake pdf content"))
        content = asyncio.run(read_file_contents(upload))

        self.assertEqual(content, "Hello PDF page 1. Hello PDF page 2.")

    def test_txt_parser_utf8(self):
        upload = UploadFile(filename="test.txt", file=BytesIO("你好 UTF-8".encode()))
        self.assertEqual(asyncio.run(read_file_contents(upload)), "你好 UTF-8")

    def test_txt_parser_gbk(self):
        upload = UploadFile(filename="test.txt", file=BytesIO("你好 GBK".encode("gbk")))
        self.assertEqual(asyncio.run(read_file_contents(upload)), "你好 GBK")

    def test_parse_failure_raises_instead_of_returning_empty_text(self):
        """解析失败必须让调用方知道。

        注册表本身不抛异常（总是返回 ParseResult），但这条上传链路的调用方
        ``/index/{name}/upload_file_by_QA`` 期望"解析不了就报错"——如果悄悄
        返回空字符串，会拿一份空内容去让 LLM 生成问答对，产出一堆凭空捏造的
        QA 塞进知识库，比直接报错危害大得多。
        """
        upload = UploadFile(filename="broken.docx", file=BytesIO(b"not a real docx"))
        with self.assertRaises(DocumentParseError) as ctx:
            asyncio.run(read_file_contents(upload))
        self.assertIn("broken.docx", str(ctx.exception))

    def test_unsupported_extension_raises(self):
        upload = UploadFile(filename="archive.zip", file=BytesIO(b"PK\x03\x04"))
        with self.assertRaises(DocumentParseError):
            asyncio.run(read_file_contents(upload))

    def test_missing_ocr_dependency_raises_distinguishable_error(self):
        """缺可选依赖和"文件本身坏了"是两回事，上层要能区分。

        ``ParserUnavailableError`` 是 ``DocumentParseError`` 的子类：不关心区别
        的调用方按基类捕获即可，想提示"装一下 OCR 依赖"的调用方可以单独捕获
        子类。这里断言的是这个层级关系没被破坏。
        """
        upload = UploadFile(filename="flow.jpg", file=BytesIO(b"\xff\xd8\xff fake jpeg"))
        try:
            asyncio.run(read_file_contents(upload))
        except ParserUnavailableError as e:
            self.assertIn("rapidocr", str(e).lower())
            self.assertIsInstance(e, DocumentParseError)
        except DocumentParseError:
            # 装了 OCR 依赖的环境：这段假字节不是合法图片，归类为 FAILURE 也
            # 是正确行为，不该让这个测试在装了 ocr extra 的机器上变红。
            pass


if __name__ == "__main__":
    unittest.main()
